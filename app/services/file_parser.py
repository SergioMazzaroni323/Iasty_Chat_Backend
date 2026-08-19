from io import BytesIO

from app.services.pdf import parse_pdf


def parse_txt(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_docx(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_file_bytes(filename: str, data: bytes) -> str:
    if not data:
        raise ValueError(f"{filename} is empty")

    lower = filename.lower()
    if lower.endswith(".txt"):
        return parse_txt(data).strip()
    if lower.endswith(".pdf"):
        text, _ = parse_pdf(data)
        return text.strip()
    if lower.endswith(".docx"):
        return parse_docx(data).strip()
    if lower.endswith(".doc"):
        raise ValueError(f"{filename}: .doc is not supported. Please use .docx")

    raise ValueError(f"{filename}: unsupported file type")


def format_file_import(filename: str, text: str) -> str:
    return f"--- From file: {filename} ---\n{text.strip()}"


def merge_content(existing: str, addition: str) -> str:
    existing = existing.strip()
    addition = addition.strip()
    if not addition:
        return existing
    if not existing:
        return addition
    return f"{existing}\n\n{addition}"
